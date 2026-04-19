"""
core/excel_parser.py
════════════════════
Routine seed data + file parsers.

Supports:
  • get_seed_routines()   — hardcoded data (15 April 2026 routine)
  • get_seed_mappings()   — teachers + course name mappings
  • parse_routine_excel() — parse .xlsx files (flat table format)
  • parse_routine_word()  — parse .docx files in official RUB format:
                            Day | Room No. | 9.00-10.10 | 10.15-11.25 |
                            11.30-12.40 | 12.40-1.35 (Prayer) | 1.35-2.45 | 2.50-4.00
                            Cell format: TEACHER_CODE (COURSE_CODE)
                            e.g.  PKP (MGT-3102)

Effective from: 15 April 2026
"""
import re
from typing import List, Dict

# ── Routine effective date ────────────────────────────────────
ROUTINE_UPDATED_DATE = "15 April 2026"

# ── Course → (program, year, semester) ───────────────────────
COURSE_META: Dict[str, tuple] = {
    # BBA 1st Year, 2nd Semester (Session: 2024-25)
    'MGT-1201': ('BBA', 1, 2),
    'GED-1202': ('BBA', 1, 2),
    'MGT-1203': ('BBA', 1, 2),
    'GED-1204': ('BBA', 1, 2),
    'GED-1205': ('BBA', 1, 2),
    'GED-1250': ('BBA', 1, 2),
    'GED-1203': ('BBA', 1, 2),
    'MGT-1202': ('BBA', 1, 2),
    'MGT-1204': ('BBA', 1, 2),
    'MGT-1205': ('BBA', 1, 2),
    # BBA 2nd Year, 1st Semester (Session: 2023-24)
    'MGT-2101': ('BBA', 2, 1),
    'GED-2102': ('BBA', 2, 1),
    'GED-2103': ('BBA', 2, 1),
    'GED-2104': ('BBA', 2, 1),
    'MGT-2105': ('BBA', 2, 1),
    # BBA 3rd Year, 1st Semester (Session: 2022-23)
    'MGT-3101': ('BBA', 3, 1),
    'MGT-3102': ('BBA', 3, 1),
    'MGT-3103': ('BBA', 3, 1),
    'MGT-3104': ('BBA', 3, 1),
    'MGT-3105': ('BBA', 3, 1),
    # BBA 3rd Year, 2nd Semester (Session: 2021-22)
    'MGT-3201': ('BBA', 3, 2),
    'MGT-3202': ('BBA', 3, 2),
    'MGT-3203': ('BBA', 3, 2),
    'MGT-3204': ('BBA', 3, 2),
    'MGT-3205': ('BBA', 3, 2),
    'GED-3203': ('BBA', 3, 2),
    # MBA 1st Semester (Session: 2024-25)
    'HRM-5101': ('MBA', 1, 1),
    'HRM-5102': ('MBA', 1, 1),
    'HRM-5103': ('MBA', 1, 1),
    'HRM-5104': ('MBA', 1, 1),
    'HRM-5105': ('MBA', 1, 1),
    # MBA 2nd Semester (Session: 2023-24)
    'HRM-5201': ('MBA', 2, 1),
    'HRM-5202': ('MBA', 2, 1),
    'HRM-5203': ('MBA', 2, 1),
    'HRM-5204': ('MBA', 2, 1),
    'HRM-5205': ('MBA', 2, 1),
}

# ── Valid days ─────────────────────────────────────────────────
VALID_DAYS = {'Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday',
              'Friday', 'Saturday'}

# ── Official display format for time_slot field ───────────────
_SLOT_DISPLAY: Dict[tuple, str] = {
    ('09:00', '10:10'): '9.00-10.10',
    ('10:15', '11:25'): '10.15-11.25',
    ('11:30', '12:40'): '11.30-12.40',
    ('13:35', '14:45'): '1.35-2.45',
    ('14:50', '16:00'): '2.50-4.00',
}

# ── Time slot column header → (time_start, time_end) ─────────
SLOT_MAP: Dict[str, tuple] = {
    '9.00-10.10':   ('09:00', '10:10'),
    '9:00-10:10':   ('09:00', '10:10'),
    '09:00-10:10':  ('09:00', '10:10'),
    '10.15-11.25':  ('10:15', '11:25'),
    '10:15-11:25':  ('10:15', '11:25'),
    '11.30-12.40':  ('11:30', '12:40'),
    '11:30-12.40':  ('11:30', '12:40'),
    '11:30-12:40':  ('11:30', '12:40'),
    '1.35-2.45':    ('13:35', '14:45'),
    '1:35-2:45':    ('13:35', '14:45'),
    '13:35-14:45':  ('13:35', '14:45'),
    '2.50-4.00':    ('14:50', '16:00'),
    '2:50-4:00':    ('14:50', '16:00'),
    '14:50-16:00':  ('14:50', '16:00'),
}

# Columns to SKIP (Prayer & Lunch etc.)
SKIP_SLOT_KEYWORDS = {'prayer', 'lunch', 'break', '12.40', '12:40', '1.35', '1:35'}

# Cell content pattern: TEACHER_CODE (COURSE_CODE)
_CELL_PATTERN = re.compile(
    r'([A-Z]{2,6})\s*\(([A-Z]{2,6}-\d{3,4}[A-Z]?)\)',
    re.IGNORECASE
)


def _resolve_slot(raw_header: str):
    """Return (time_start, time_end) for a column header, or None to skip."""
    cleaned = raw_header.strip().replace('\n', ' ').strip()
    low = cleaned.lower()
    for kw in SKIP_SLOT_KEYWORDS:
        if kw in low:
            return None
    for key, val in SLOT_MAP.items():
        if key.lower() in low or low in key.lower():
            return val
    times = re.findall(r'(\d{1,2})[.:](\d{2})', cleaned)
    if len(times) >= 2:
        h1, m1 = int(times[0][0]), int(times[0][1])
        h2, m2 = int(times[1][0]), int(times[1][1])
        if h1 < 7:
            h1 += 12
        if h2 < h1:
            h2 += 12
        return (f"{h1:02d}:{m1:02d}", f"{h2:02d}:{m2:02d}")
    return None


def _parse_cell(text: str):
    """Parse 'PKP (MGT-3102)' → ('PKP', 'MGT-3102') or None."""
    text = text.strip().replace('\n', ' ')
    if not text:
        return None
    m = _CELL_PATTERN.search(text)
    if m:
        return m.group(1).upper(), m.group(2).upper()
    return None


def _build_entry(day: str, room: str, time_start: str, time_end: str,
                 teacher_code: str, course_code: str,
                 session: str = '2025-26') -> Dict:
    meta = COURSE_META.get(course_code.upper(), ('ALL', 0, 0))
    time_slot = _SLOT_DISPLAY.get((time_start, time_end),
                                   f"{time_start}-{time_end}")
    return {
        'day':             day,
        'room_no':         room,
        'time_slot':       time_slot,
        'time_start':      time_start,
        'time_end':        time_end,
        'course_code':     course_code.upper(),
        'teacher_code':    teacher_code.upper(),
        'program':         meta[0],
        'course_year':     meta[1],
        'course_semester': meta[2],
        'session':         session,
    }


# ══════════════════════════════════════════════════════════════
# WORD (.docx) PARSER — official RUB routine format
# ══════════════════════════════════════════════════════════════

def parse_routine_word(file_path: str) -> List[Dict]:
    """
    Parse the official Rabindra University Word routine table.

    Expected Word table columns:
      Day | Room No. | 9.00-10.10 | 10.15-11.25 | 11.30-12.40 |
      12.40-1.35 (Prayer) | 1.35-2.45 | 2.50-4.00

    Cell content: TEACHER_CODE (COURSE_CODE)  e.g. PKP (MGT-3102)
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document(file_path)
    if not doc.tables:
        raise ValueError("No tables found in Word document.")

    entries: List[Dict] = []

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        header_cells = [cell.text for cell in rows[0].cells]
        col_slots: Dict[int, tuple] = {}
        day_col  = None
        room_col = None

        for i, hdr in enumerate(header_cells):
            low = hdr.strip().lower()
            if low in ('day', 'days', ''):
                day_col = i
                col_slots[i] = None
            elif low in ('room', 'room no', 'room no.', 'room number', 'room_no'):
                room_col = i
                col_slots[i] = None
            else:
                col_slots[i] = _resolve_slot(hdr)

        if day_col is None:
            day_col = 0
        if room_col is None:
            room_col = 1

        current_day = ''
        for row in rows[1:]:
            cells = row.cells
            raw_day = cells[day_col].text.strip() if day_col < len(cells) else ''
            if raw_day:
                candidate = raw_day.strip().title()
                if candidate in VALID_DAYS:
                    current_day = candidate
            if not current_day:
                continue

            room = cells[room_col].text.strip() if room_col < len(cells) else ''
            if not room:
                continue

            for col_idx, slot in col_slots.items():
                if slot is None:
                    continue
                if col_idx >= len(cells):
                    continue
                parsed = _parse_cell(cells[col_idx].text)
                if parsed is None:
                    continue
                teacher_code, course_code = parsed
                time_start, time_end = slot
                entries.append(_build_entry(
                    day=current_day, room=room,
                    time_start=time_start, time_end=time_end,
                    teacher_code=teacher_code, course_code=course_code,
                ))

    if not entries:
        raise ValueError(
            "No valid routine entries found. Make sure the Word file uses "
            "the official RUB format: Day | Room No. | time slots, "
            "with cells like 'PKP (MGT-3102)'."
        )

    seen: set = set()
    unique: List[Dict] = []
    for e in entries:
        key = (e['day'], e['room_no'], e['time_start'], e['course_code'])
        if key not in seen:
            seen.add(key)
            unique.append(e)
    return unique


# ══════════════════════════════════════════════════════════════
# EXCEL (.xlsx) PARSER — flat table format (legacy)
# ══════════════════════════════════════════════════════════════

def parse_routine_excel(file_path: str) -> List[Dict]:
    """
    Parse a flat .xlsx routine file.
    Required columns: day, room_no, time_start, time_end,
                      course_code, teacher_code
    """
    if file_path.lower().endswith('.docx'):
        return parse_routine_word(file_path)

    import pandas as pd
    try:
        df = pd.read_excel(file_path)
    except Exception as e:
        raise ValueError(f"Could not read Excel file: {e}")

    required = {'day', 'room_no', 'time_start', 'time_end', 'course_code', 'teacher_code'}
    cols = {c.strip().lower() for c in df.columns}
    if not required.issubset(cols):
        raise ValueError(f"Missing columns: {required - cols}")

    df.columns = [c.strip().lower() for c in df.columns]
    entries: List[Dict] = []

    for _, row in df.iterrows():
        try:
            def norm(t: str) -> str:
                p = str(t).strip().split(':')
                return f"{p[0].zfill(2)}:{p[1].zfill(2)}" if len(p) == 2 else str(t)

            entry = _build_entry(
                day          = str(row.get('day', '')).strip(),
                room         = str(row.get('room_no', '')).strip(),
                time_start   = norm(str(row.get('time_start', ''))),
                time_end     = norm(str(row.get('time_end', ''))),
                teacher_code = str(row.get('teacher_code', '')).strip().upper(),
                course_code  = str(row.get('course_code', '')).strip().upper(),
                session      = str(row.get('session', '2025-26')).strip(),
            )
            if entry['day'] and entry['course_code']:
                entries.append(entry)
        except Exception:
            continue
    return entries


# ══════════════════════════════════════════════════════════════
# SEED DATA — official routine effective 15 April 2026
# Dept. of Management, Rabindra University, Bangladesh
# Total: 60 class slots across 3 rooms, 5 days
# ══════════════════════════════════════════════════════════════

def get_seed_routines() -> List[Dict]:
    """
    Return all 60 class slots from the official routine.
    Format: (Day, Room, TimeStart, TimeEnd, TeacherCode, CourseCode)

    Time slot mapping:
      09:00-10:10 → 9.00-10.10   (Column 1)
      10:15-11:25 → 10.15-11.25  (Column 2)
      11:30-12:40 → 11.30-12.40  (Column 3)
      ── Prayer & Lunch ──────── (Skipped)
      13:35-14:45 → 1.35-2.45   (Column 4)
      14:50-16:00 → 2.50-4.00   (Column 5)
    """
    raw = [
        # ══ SUNDAY ═══════════════════════════════════════════
        # Room 101 — 4 slots
        ("Sunday", "101",  "10:15", "11:25", "PKP", "MGT-3102"),
        ("Sunday", "101",  "11:30", "12:40", "PKP", "GED-2102"),
        ("Sunday", "101",  "13:35", "14:45", "KHR", "GED-2103"),
        ("Sunday", "101",  "14:50", "16:00", "HR",  "GED-3203"),
        # Room 201 — 3 slots
        ("Sunday", "201",  "10:15", "11:25", "THT", "MGT-2105"),
        ("Sunday", "201",  "13:35", "14:45", "AH",  "MGT-3201"),
        ("Sunday", "201",  "14:50", "16:00", "AH",  "HRM-5205"),
        # Room 1001 — 3 slots
        ("Sunday", "1001", "11:30", "12:40", "AH",  "GED-1204"),
        ("Sunday", "1001", "13:35", "14:45", "PKP", "MGT-1201"),
        ("Sunday", "1001", "14:50", "16:00", "MT",  "GED-1205"),

        # ══ MONDAY ═══════════════════════════════════════════
        # Room 101 — 4 slots
        ("Monday", "101",  "09:00", "10:10", "PKP", "GED-2102"),
        ("Monday", "101",  "10:15", "11:25", "KHR", "HRM-5105"),
        ("Monday", "101",  "11:30", "12:40", "THT", "HRM-5101"),
        ("Monday", "101",  "13:35", "14:45", "PKP", "HRM-5102"),
        # Room 201 — 4 slots
        ("Monday", "201",  "09:00", "10:10", "AH",  "HRM-5205"),
        ("Monday", "201",  "10:15", "11:25", "AH",  "GED-2104"),
        ("Monday", "201",  "11:30", "12:40", "HR",  "MGT-2101"),
        ("Monday", "201",  "13:35", "14:45", "MK",  "MGT-3204"),
        # Room 1001 — 4 slots
        ("Monday", "1001", "10:15", "11:25", "PKP", "HRM-5203"),
        ("Monday", "1001", "11:30", "12:40", "KHR", "HRM-5202"),
        ("Monday", "1001", "13:35", "14:45", "KHR", "MGT-1203"),
        ("Monday", "1001", "14:50", "16:00", "TI",  "GED-1202"),

        # ══ TUESDAY ══════════════════════════════════════════
        # Room 101 — 5 slots
        ("Tuesday", "101",  "09:00", "10:10", "MK",  "HRM-5103"),
        ("Tuesday", "101",  "10:15", "11:25", "PKP", "HRM-5102"),
        ("Tuesday", "101",  "11:30", "12:40", "THT", "MGT-2105"),
        ("Tuesday", "101",  "13:35", "14:45", "PKP", "MGT-3102"),
        ("Tuesday", "101",  "14:50", "16:00", "AH",  "MGT-3105"),
        # Room 201 — 5 slots
        ("Tuesday", "201",  "09:00", "10:10", "FA",  "MGT-3205"),
        ("Tuesday", "201",  "10:15", "11:25", "HR",  "MGT-3104"),
        ("Tuesday", "201",  "11:30", "12:40", "KHR", "MGT-3202"),
        ("Tuesday", "201",  "13:35", "14:45", "AH",  "GED-2104"),
        ("Tuesday", "201",  "14:50", "16:00", "HR",  "MGT-2101"),
        # Room 1001 — 3 slots
        ("Tuesday", "1001", "09:00", "10:10", "PKP", "HRM-5203"),
        ("Tuesday", "1001", "10:15", "11:25", "KHR", "HRM-5202"),
        ("Tuesday", "1001", "11:30", "12:40", "HR",  "HRM-5204"),

        # ══ WEDNESDAY ════════════════════════════════════════
        # Room 101 — 4 slots
        ("Wednesday", "101",  "10:15", "11:25", "HR",  "HRM-5204"),
        ("Wednesday", "101",  "11:30", "12:40", "HR",  "GED-3203"),
        ("Wednesday", "101",  "13:35", "14:45", "HR",  "MGT-3104"),
        ("Wednesday", "101",  "14:50", "16:00", "THT", "MGT-3101"),
        # Room 201 — 4 slots
        ("Wednesday", "201",  "09:00", "10:10", "FA",  "MGT-3205"),
        ("Wednesday", "201",  "10:15", "11:25", "MK",  "MGT-3204"),
        ("Wednesday", "201",  "11:30", "12:40", "MK",  "HRM-5201"),
        ("Wednesday", "201",  "13:35", "14:45", "MK",  "HRM-5201"),
        # Room 1001 — 5 slots
        ("Wednesday", "1001", "09:00", "10:10", "AH",  "HRM-5104"),
        ("Wednesday", "1001", "10:15", "11:25", "KHR", "HRM-5105"),
        ("Wednesday", "1001", "11:30", "12:40", "THT", "HRM-5101"),
        ("Wednesday", "1001", "13:35", "14:45", "PKP", "MGT-1201"),
        ("Wednesday", "1001", "14:50", "16:00", "MT",  "GED-1250"),

        # ══ THURSDAY ═════════════════════════════════════════
        # Room 101 — 5 slots
        ("Thursday", "101",  "09:00", "10:10", "MK",  "HRM-5103"),
        ("Thursday", "101",  "10:15", "11:25", "KHR", "MGT-3202"),
        ("Thursday", "101",  "11:30", "12:40", "KHR", "MGT-3103"),
        ("Thursday", "101",  "13:35", "14:45", "AH",  "MGT-3105"),
        ("Thursday", "101",  "14:50", "16:00", "THT", "MGT-3101"),
        # Room 201 — 4 slots
        ("Thursday", "201",  "09:00", "10:10", "FA",  "MGT-3205"),
        ("Thursday", "201",  "10:15", "11:25", "AH",  "HRM-5104"),
        ("Thursday", "201",  "11:30", "12:40", "AH",  "MGT-3201"),
        ("Thursday", "201",  "13:35", "14:45", "KHR", "GED-2103"),
        # Room 1001 — 3 slots
        ("Thursday", "1001", "09:00", "10:10", "KHR", "MGT-1203"),
        ("Thursday", "1001", "10:15", "11:25", "MT",  "GED-1250"),
        ("Thursday", "1001", "11:30", "12:40", "TI",  "GED-1202"),
    ]

    return [
        _build_entry(
            day=r[0], room=r[1],
            time_start=r[2], time_end=r[3],
            teacher_code=r[4], course_code=r[5],
        )
        for r in raw
    ]


def get_seed_mappings() -> List[Dict]:
    """Return all teacher + course mappings for Dept. of Management."""
    data = [
        # ── Teachers ──────────────────────────────────────────
        ("FA",       "Professor Dr. Feroz Ahmed",                  "teacher"),
        ("HR",       "Habibur Rahaman",                            "teacher"),
        ("MK",       "Malina Khatun",                              "teacher"),
        ("PKP",      "Proshanta Kumar Podder",                     "teacher"),
        ("TI",       "Md. Tarequl Islam",                          "teacher"),
        ("AH",       "Alamgir Hossain",                            "teacher"),
        ("KHR",      "Md. Kazi Hafizur Rahman",                    "teacher"),
        ("THT",      "Tamima Hasan Taishi",                        "teacher"),
        ("MT",       "Mst. Mohsina Tanjim",                        "teacher"),
        # ── BBA 1st Year 2nd Semester ─────────────────────────
        ("MGT-1201", "Business Communication",                     "course"),
        ("GED-1202", "Computer in Business",                       "course"),
        ("MGT-1203", "Principles of Finance",                      "course"),
        ("GED-1204", "Microeconomics",                             "course"),
        ("GED-1205", "Communicative English",                      "course"),
        ("GED-1250", "Communicative English II",                   "course"),
        # ── BBA 2nd Year 1st Semester ─────────────────────────
        ("MGT-2101", "Organization Behavior",                      "course"),
        ("GED-2102", "Statistics-I",                               "course"),
        ("GED-2103", "Commercial Law",                             "course"),
        ("GED-2104", "Macro Economics",                            "course"),
        ("MGT-2105", "Financial Management",                       "course"),
        # ── BBA 3rd Year 1st Semester ─────────────────────────
        ("MGT-3101", "Bank Management",                            "course"),
        ("MGT-3102", "Industrial Law",                             "course"),
        ("MGT-3103", "Insurance and Risk Management",              "course"),
        ("MGT-3104", "Taxation & Auditing",                        "course"),
        ("MGT-3105", "Innovation & Change Management",             "course"),
        # ── BBA 3rd Year 2nd Semester ─────────────────────────
        ("MGT-3201", "Industrial Relations",                       "course"),
        ("MGT-3202", "Management Information System",              "course"),
        ("MGT-3203", "Business Ethics and CSR",                    "course"),
        ("MGT-3204", "Management Science",                         "course"),
        ("MGT-3205", "Entrepreneurship and SME Management",        "course"),
        ("GED-3203", "Business Ethics and CSR",                    "course"),
        # ── MBA 1st Semester ──────────────────────────────────
        ("HRM-5101", "Human Resource Planning and Policy",         "course"),
        ("HRM-5102", "Training and Development",                   "course"),
        ("HRM-5103", "Compensation Management",                    "course"),
        ("HRM-5104", "Conflict Management",                        "course"),
        ("HRM-5105", "Strategic Human Resource Management",        "course"),
        # ── MBA 2nd Semester ──────────────────────────────────
        ("HRM-5201", "International Human Resource Management",    "course"),
        ("HRM-5202", "Human Resource Information System",          "course"),
        ("HRM-5203", "Career Planning and Development",            "course"),
        ("HRM-5204", "Performance Management",                     "course"),
        ("HRM-5205", "Contemporary Issues in HRM",                 "course"),
    ]
    return [{"code": d[0], "full_name": d[1], "type": d[2]} for d in data]